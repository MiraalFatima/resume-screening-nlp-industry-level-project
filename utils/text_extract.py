
# File: utils/text_extract.py

from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document
import io

def extract_text_from_file(file_path_or_buffer):
    """
    Extracts text from a .pdf or .docx file, accepting either a file path or a file-like object.
    """
    # Check if the input is a path (string) or a file-like object
    if isinstance(file_path_or_buffer, str):
        file_extension = file_path_or_buffer.split('.')[-1].lower()
        if file_extension == "pdf":
            return extract_pdf_text(file_path_or_buffer)
        elif file_extension == "docx":
            doc = Document(file_path_or_buffer)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            raise ValueError("Unsupported file format. Only .pdf and .docx are supported.")
    # Handle file-like objects (e.g., from Streamlit's file_uploader)
    else:
        if file_path_or_buffer.name.endswith(".pdf"):
            # pdfminer.six needs bytes, so we read the buffer
            return extract_pdf_text(io.BytesIO(file_path_or_buffer.read()))
        elif file_path_or_buffer.name.endswith(".docx"):
            doc = Document(file_path_or_buffer)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            raise ValueError("Unsupported file format. Only .pdf and .docx are supported.")